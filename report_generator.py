import io
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


class ReportGenerator:
    """Generate professional DOCX reports with analysis and visualizations"""
    
    def __init__(self, df, analysis_data, output_path):
        """
        Initialize report generator
        
        Args:
            df: Original DataFrame
            analysis_data: Analysis results from DataAnalyzer
            output_path: Path to save the DOCX report
        """
        self.df = df
        self.analysis = analysis_data
        self.output_path = output_path
        self.doc = Document()
        
        # Set up styling
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup document styles"""
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
    
    def _add_heading(self, text, level=1):
        """Add heading to document"""
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_paragraph(self, text, bold=False, size=11):
        """Add paragraph to document"""
        p = self.doc.add_paragraph(text)
        if bold:
            for run in p.runs:
                run.font.bold = True
        for run in p.runs:
            run.font.size = Pt(size)
    
    def _format_value(self, value):
        """Format numeric values for display"""
        if isinstance(value, float):
            return f"{value:.4f}"
        return str(value)
    
    def _add_table(self, data_dict, title=""):
        """Add table to document"""
        if not data_dict:
            return
        
        rows = len(data_dict) + 1
        cols = 2
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = "Metric"
        header_cells[1].text = "Value"
        
        for i, (key, value) in enumerate(data_dict.items(), 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(key)
            row_cells[1].text = self._format_value(value)
    
    def _save_figure_to_bytes(self, fig):
        """Convert matplotlib figure to bytes"""
        img_stream = io.BytesIO()
        fig.savefig(img_stream, format='png', dpi=100, bbox_inches='tight')
        img_stream.seek(0)
        plt.close(fig)
        return img_stream
    
    def _create_numeric_distribution_plots(self):
        """Create distribution plots for numeric columns"""
        numeric_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()
        
        for col in numeric_cols[:6]:  # Limit to first 6 columns
            fig, axes = plt.subplots(1, 2, figsize=(12, 4))
            
            # Histogram
            axes[0].hist(self.df[col].dropna(), bins=30, color='steelblue', edgecolor='black', alpha=0.7)
            axes[0].set_title(f'Distribution of {col}', fontsize=12, fontweight='bold')
            axes[0].set_xlabel(col)
            axes[0].set_ylabel('Frequency')
            axes[0].grid(True, alpha=0.3)
            
            # Box plot
            axes[1].boxplot(self.df[col].dropna())
            axes[1].set_title(f'Box Plot of {col}', fontsize=12, fontweight='bold')
            axes[1].set_ylabel(col)
            axes[1].grid(True, alpha=0.3)
            
            img_stream = self._save_figure_to_bytes(fig)
            self.doc.add_picture(img_stream, width=Inches(6.0))
            self.doc.add_paragraph()
    
    def _create_correlation_heatmap(self):
        """Create correlation heatmap"""
        numeric_cols = self.df.select_dtypes(include=[np.number]).columns.tolist()
        
        if len(numeric_cols) < 2:
            return
        
        fig, ax = plt.subplots(figsize=(10, 8))
        corr_matrix = self.df[numeric_cols].corr()
        
        sns.heatmap(corr_matrix, annot=True, fmt='.2f', cmap='coolwarm', center=0, 
                   square=True, ax=ax, cbar_kws={"shrink": 0.8})
        ax.set_title('Correlation Matrix Heatmap', fontsize=14, fontweight='bold', pad=20)
        
        img_stream = self._save_figure_to_bytes(fig)
        self.doc.add_picture(img_stream, width=Inches(6.0))
        self.doc.add_paragraph()
    
    def _create_categorical_plots(self):
        """Create visualizations for categorical columns"""
        categorical_cols = self.df.select_dtypes(include=['object', 'category']).columns.tolist()
        
        for col in categorical_cols[:4]:  # Limit to first 4 columns
            value_counts = self.df[col].value_counts()
            
            if len(value_counts) > 10:
                value_counts = value_counts.head(10)
            
            fig, ax = plt.subplots(figsize=(10, 5))
            value_counts.plot(kind='bar', ax=ax, color='steelblue', edgecolor='black', alpha=0.7)
            ax.set_title(f'Value Counts: {col}', fontsize=12, fontweight='bold')
            ax.set_xlabel(col)
            ax.set_ylabel('Count')
            ax.grid(True, alpha=0.3, axis='y')
            plt.xticks(rotation=45, ha='right')
            
            img_stream = self._save_figure_to_bytes(fig)
            self.doc.add_picture(img_stream, width=Inches(6.0))
            self.doc.add_paragraph()
    
    def _create_missing_data_visualization(self):
        """Create visualization of missing data"""
        missing_data = self.df.isnull().sum()
        missing_data = missing_data[missing_data > 0]
        
        if missing_data.empty:
            return
        
        fig, ax = plt.subplots(figsize=(10, 5))
        missing_data.plot(kind='bar', ax=ax, color='coral', edgecolor='black', alpha=0.7)
        ax.set_title('Missing Values by Column', fontsize=12, fontweight='bold')
        ax.set_xlabel('Column')
        ax.set_ylabel('Count of Missing Values')
        ax.grid(True, alpha=0.3, axis='y')
        plt.xticks(rotation=45, ha='right')
        
        img_stream = self._save_figure_to_bytes(fig)
        self.doc.add_picture(img_stream, width=Inches(6.0))
        self.doc.add_paragraph()
    
    def generate_report(self):
        """Generate complete report"""
        # Title Page
        title = self.doc.add_heading('Data Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        subtitle = self.doc.add_paragraph('Comprehensive Data Analysis & Profiling')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0]
        subtitle_format.font.size = Pt(14)
        
        self.doc.add_paragraph()
        date_para = self.doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()
        
        # Table of Contents
        self._add_heading('Table of Contents', 1)
        toc_items = [
            '1. Executive Summary',
            '2. Data Overview',
            '3. Missing Data Analysis',
            '4. Numeric Columns Statistics',
            '5. Distribution Analysis',
            '6. Categorical Columns Analysis',
            '7. Outlier Detection',
            '8. Class Balance Analysis',
            '9. Correlation Analysis',
            '10. Visualizations',
        ]
        for item in toc_items:
            self.doc.add_paragraph(item, style='List Bullet')
        
        self.doc.add_page_break()
        
        # 1. Executive Summary
        self._add_heading('1. Executive Summary', 1)
        overview = self.analysis['overview']
        self._add_paragraph(f"Total Rows: {overview['total_rows']}")
        self._add_paragraph(f"Total Columns: {overview['total_columns']}")
        self._add_paragraph(f"Numeric Columns: {overview['numeric_columns']}")
        self._add_paragraph(f"Categorical Columns: {overview['categorical_columns']}")
        self._add_paragraph(f"Duplicate Rows: {overview['duplicate_rows']} ({overview['duplicate_percentage']:.2f}%)")
        self._add_paragraph(f"Memory Usage: {overview['memory_usage_mb']:.2f} MB")
        
        self.doc.add_page_break()
        
        # 2. Data Overview
        self._add_heading('2. Data Overview', 1)
        self._add_paragraph('Column Information:')
        col_data = {
            'Column': self.df.columns.tolist(),
            'Data Type': [str(dtype) for dtype in self.df.dtypes.tolist()]
        }
        col_df = pd.DataFrame(col_data)
        rows = len(col_df) + 1
        cols = 2
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        header_cells[0].text = "Column"
        header_cells[1].text = "Data Type"
        for i, row in col_df.iterrows():
            row_cells = table.rows[i+1].cells
            row_cells[0].text = str(row['Column'])
            row_cells[1].text = str(row['Data Type'])
        
        self.doc.add_page_break()
        
        # 3. Missing Data Analysis
        self._add_heading('3. Missing Data Analysis', 1)
        missing = self.analysis['missing_data']
        self._add_paragraph(f"Total Missing Cells: {missing['total_missing_cells']}")
        self._add_paragraph(f"Total Cells: {missing['total_cells']}")
        self._add_paragraph(f"Missing Percentage: {missing['missing_cell_percentage']:.2f}%")
        
        if missing['total_missing_cells'] > 0:
            self._add_paragraph('Missing Values by Column:', bold=True)
            missing_table_data = {}
            for col, count in missing['missing_count'].items():
                if count > 0:
                    missing_table_data[col] = f"{count} ({missing['missing_percentage'][col]:.2f}%)"
            self._add_table(missing_table_data)
            
            self._add_heading('Missing Data Visualization', 2)
            self._create_missing_data_visualization()
        else:
            self._add_paragraph("No missing data detected in this dataset.")
        
        self.doc.add_page_break()
        
        # 4. Numeric Columns Statistics
        self._add_heading('4. Numeric Columns Statistics', 1)
        numeric_stats = self.analysis['numeric_statistics']
        
        for col, stats_data in numeric_stats.items():
            self._add_heading(f"Column: {col}", 2)
            stats_table_data = {
                'Count': stats_data['count'],
                'Mean': self._format_value(stats_data['mean']),
                'Median': self._format_value(stats_data['median']),
                'Std Dev': self._format_value(stats_data['std_dev']),
                'Min': self._format_value(stats_data['min']),
                'Max': self._format_value(stats_data['max']),
                'Q25': self._format_value(stats_data['q25']),
                'Q75': self._format_value(stats_data['q75']),
                'IQR': self._format_value(stats_data['iqr']),
                'Range': self._format_value(stats_data['range']),
                'Coefficient of Variation': f"{stats_data['coefficient_of_variation']:.2f}%",
            }
            self._add_table(stats_table_data)
            self.doc.add_paragraph()
        
        self.doc.add_page_break()
        
        # 5. Distribution Analysis
        self._add_heading('5. Distribution Analysis', 1)
        dist_analysis = self.analysis['distribution']
        
        for col, dist_data in dist_analysis.items():
            self._add_heading(f"Column: {col}", 2)
            dist_table = {
                'Is Normal': 'Yes' if dist_data['is_normal'] else 'No',
                'Normality P-Value': f"{dist_data['normality_p_value']:.6f}",
                'Skewness': self._format_value(dist_data['skewness']),
                'Skewness Type': dist_data['skew_type'],
                'Kurtosis': self._format_value(dist_data['kurtosis']),
                'Kurtosis Type': dist_data['kurt_type'],
            }
            self._add_table(dist_table)
            self.doc.add_paragraph()
        
        self.doc.add_page_break()
        
        # 6. Categorical Columns Analysis
        self._add_heading('6. Categorical Columns Analysis', 1)
        categorical = self.analysis['categorical']
        
        for col, cat_data in categorical.items():
            self._add_heading(f"Column: {col}", 2)
            cat_table = {
                'Unique Values': cat_data['unique_values'],
                'Top Value': cat_data['top_value'],
                'Top Value Count': cat_data['top_value_count'],
                'Missing Values': cat_data['missing_count'],
            }
            self._add_table(cat_table)
            self.doc.add_paragraph()
        
        self.doc.add_page_break()
        
        # 7. Outlier Detection
        self._add_heading('7. Outlier Detection', 1)
        outliers = self.analysis['outliers']
        
        for col, outlier_data in outliers.items():
            self._add_heading(f"Column: {col}", 2)
            outlier_table = {
                'Lower Bound': self._format_value(outlier_data['lower_bound']),
                'Upper Bound': self._format_value(outlier_data['upper_bound']),
                'Outlier Count': outlier_data['outlier_count'],
                'Outlier Percentage': f"{outlier_data['outlier_percentage']:.2f}%",
            }
            self._add_table(outlier_table)
            self.doc.add_paragraph()
        
        self.doc.add_page_break()
        
        # 8. Class Balance Analysis
        self._add_heading('8. Class Balance Analysis', 1)
        balance = self.analysis['class_balance']
        
        if balance:
            for col, balance_data in balance.items():
                self._add_heading(f"Column: {col}", 2)
                balance_table = {
                    'Unique Classes': balance_data['unique_classes'],
                    'Is Balanced': 'Yes' if balance_data['is_balanced'] else 'No',
                    'Imbalance Ratio': f"{balance_data['imbalance_ratio']:.2f}",
                }
                self._add_table(balance_table)
                
                self._add_paragraph('Class Distribution:', bold=True)
                dist_table_data = {f"{k}": f"{v:.2f}%" for k, v in balance_data['class_distribution'].items()}
                self._add_table(dist_table_data)
                self.doc.add_paragraph()
        else:
            self._add_paragraph("No categorical columns to analyze for class balance.")
        
        self.doc.add_page_break()
        
        # 9. Correlation Analysis
        self._add_heading('9. Correlation Analysis', 1)
        correlation = self.analysis['correlation']
        
        if correlation:
            if correlation['strong_correlations']:
                self._add_paragraph('Strong Correlations (abs > 0.7):', bold=True)
                strong_corr_table = {k: f"{v:.4f}" for k, v in correlation['strong_correlations'].items()}
                self._add_table(strong_corr_table)
            else:
                self._add_paragraph("No strong correlations (abs > 0.7) detected.")
        
        self.doc.add_page_break()
        
        # 10. Visualizations
        self._add_heading('10. Visualizations', 1)
        
        self._add_heading('Distribution Plots', 2)
        self._create_numeric_distribution_plots()
        
        self._add_heading('Categorical Analysis', 2)
        self._create_categorical_plots()
        
        self._add_heading('Correlation Matrix', 2)
        self._create_correlation_heatmap()
        
        # Footer
        self.doc.add_page_break()
        self._add_heading('Report Summary', 1)
        self._add_paragraph("This report provides a comprehensive analysis of the provided dataset, including:")
        footer_points = [
            "Data structure and composition overview",
            "Missing value analysis and patterns",
            "Statistical summary of numeric columns",
            "Distribution characteristics and normality tests",
            "Categorical variable analysis",
            "Outlier detection using IQR method",
            "Class balance assessment for categorical variables",
            "Correlation analysis between numeric variables",
            "Visual representations of data distributions and relationships",
        ]
        for point in footer_points:
            self.doc.add_paragraph(point, style='List Bullet')
        
        self.doc.add_paragraph()
        self.doc.add_paragraph(f"Report generated using Data Analysis Tool - {datetime.now().strftime('%Y-%m-%d')}")
        
        # Save document
        self.doc.save(self.output_path)
        print(f"Report saved to: {self.output_path}")
